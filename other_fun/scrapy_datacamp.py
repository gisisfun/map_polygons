#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Apr 23 21:58:02 2020
@author: me

DataCamp Summary of Courses Tracks and Projects
Sourced from html from profile page (my_profile.txt) and pdf certificates 
sotored in subdirectories (Courses, Skills, Career)
"""
import pdftotext
import os
import numpy as np
import pandas as pd
from scrapy import Selector
import matplotlib.pyplot as plt
import re

from docx import Document


document = Document()
paragraphs = document.paragraphs

from docx.shared import Pt

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

document.styles['Normal']

def extract(page_text,c_type):
    ref=re.findall('#\d+.\d+',page_text)
    page_lines = page_text.split('\n')
    if c_type=='Courses':        
        if len(page_lines[5]) < 1:
            course = page_lines[4]
        else:
            course = page_lines[4]+' '+page_lines[5]
        date = re.findall('[A-Z][a-z]{2}\s\d{2},\s\d{4}',page_text)[0]
    else:
        course = page_lines[2]
        date=np.nan
    return ref,course,date

def find_cert_pdf_files(path_to_files):
    list_out=[]
    for (root,dirs,files) in os.walk(path_to_files, topdown=True):
        print(dirs,root)
        dirs_list=dirs
        for dir_name in dirs_list:
            the_path=path_to_files +'/'+dir_name
            for (root,dirs,files) in os.walk(the_path, topdown=True):
                #print("bbbb",path_to_files,files)
                if 'certificate.pdf' in files:
                    
                    the_file=''+root+'/certificate.pdf'
                    # Load your PDF
                    with open(the_file, "rb") as f:
                        pdf = pdftotext.PDF(f)

                        # Iterate over all the pages
                        for page in pdf:
                            #print(page)
                            page_text= page
                            (ref, course, date) = extract(page_text, dir_name)
                            list_out.append([ref[0], course, date, dir_name])
        return list_out

def doctable(data, tabletitle):
    data = pd.DataFrame(data)  # My input data is in the 2D list form
    document.add_heading(tabletitle)
    table = document.add_table(rows=1, cols=2)
    #not working
    #hdr_cells = table.rows[0].cells
    #for i,val in enumerate(data.columns):
    #    print(i)
        #hdr_cells[i].text = val
    table = document.add_table(rows=(data.shape[0]), cols=data.shape[1])  # First row are table headers!
    #document.styles("Normal")
    for i, column in enumerate(data) :
        for row in range(data.shape[0]) :
            table.cell(row, i).line_spacing = 0.5
            table.cell(row, i).text = str(data[column][row])

def print_me(text='',to_docx=True, style='Normal'):
    if (to_docx):
        document.add_paragraph(text)
    print(text)


def just_words(raw_html,func,non_extra=[]):
    '''
    Finds all words in html 
    
    input:
        raw_html - list of strings returned from scrapy query
        
    returns:
        list of strings without html text
    '''
    non_words = ['block__main', 'class', 
                 'dc', 'div', 'h4', 'mb',
                 'mt', 'p', 'track', 'u',
                 'course','block__technology'] + non_extra
    final_list=[]
    for y in raw_html:
        words_l=[]
        html_words =re.findall('[\w]+',y)
        for x in html_words:
            if x not in non_words:
                words_l.append(func(x))
        final_list.append(' '.join(words_l))
    return final_list

def nothing(value):
    return value

def is_amp(f_value):
    output=f_value
    if f_value=="amp":
        output='&'   
    return output

if __name__ == "__main__":
    #course_as = sel.css( 'div.course-block > a' )
    # Selecting all href attributes chaining with css
    #hrefs_from_css = course_as.css( '::attr(href)' )
    # Selecting all href attributes chaining with xpath
    #hrefs_from_xpath = course_as.xpath( './@href' )
    css_dc_counts = 'strong.dc-u-m-none::text'

    xsel_course_lang = '//div[contains(@class,"course-block__technology course-block__technology--")]'

    css_course_list = 'h4.course-block__title::text'
    css_course_urls = 'a.course-block__link::attr(href)'
    css_course_descriptions = 'p.course-block__description::text'
    
    css_per_topic_names = 'h4.dc-u-mb-4::text' 
    css_per_topic_data = 'p.dc-u-mt-0::text' #first 12
    css_other_track = 'span.dc-dropdown--nav__track-name::text'
    css_topic_urls = 'section.profile-topics a.shim::attr(href)'
    
    css_track_names ='div.track-block__main'
    css_track_urls = 'section.profile-tracks a.shim::attr(href)'
    css_track_comp = 'div div h5::text'
    #css_track_names = 'div.track-block__main h4::text'
    css_urls = 'a::attr(href)'

    css_titles_text = 'title::text'
    css_project_urls='section.profile-courses a.shim::attr(href)'
    css_project_names='section.profile-courses h5.dc-project-block__title::text'
    
    thefile = open('my_profile.txt','r')

    print('getting pdf cert data')
    cert_list = find_cert_pdf_files('Datacamp')
                
    cert_df = pd.DataFrame(cert_list, columns=['ref','course','date','c_type'])
    #cert_df.ref = cert_df.ref.str.strip()
    cert_df.course = cert_df.course\
                     .str.replace('COMPLETED ON','')\
                     .str.strip()\
                     .str.replace(' Track','')
    cert_df['date_fmt'] = pd.to_datetime(cert_df.date)
    cert_df['course'] = cert_df.course.str.replace('Ef cient','Efficient')
    cert_df['course_lower'] = cert_df.course.str.lower()
    

    html = thefile.read()
    sel = Selector(text=html)
    the_title = sel.css(css_titles_text).extract_first()
    document.add_heading(the_title, 0)
    print(the_title)
    print('')

    document.add_heading('XP Points by Topic', level=1)
    print('XP by topic')

    dc_counts = sel.css(css_dc_counts).extract()
    print_me('Total XP '+dc_counts[0]+' Total Courses '+dc_counts[1] + 
             ' Total Exercises '+dc_counts[2])
    cert_counts = cert_df.c_type.value_counts()
    print_me('Skill Tracks '+str(cert_counts[1])+' Career Tracks '+str(cert_counts[2])) 
    #scape web page for content
    dc_topic_names = sel.css(css_per_topic_names).extract()
    dc_topic_data = sel.css(css_per_topic_data).extract()[:12]
    dc_topic_data = just_words(dc_topic_data,nothing,['XP'])
    # clean tex and convert to integer
    dc_topic_data = [int(x) for x in dc_topic_data]

    my_topic_xp = pd.DataFrame(list(zip(dc_topic_names,dc_topic_data)), \
                               columns =['Topic','XP'])
    # make category
    my_topic_xp.Topic = my_topic_xp.Topic.astype("category")
    my_topic_xp.sort_values(by='XP',ascending=False,inplace=True)

    plt.bar(x=my_topic_xp.Topic,height=my_topic_xp.XP)
    plt.xticks(rotation=90)
    plt.ylabel("XP")
    plt.xlabel("Topic")
    plt.title("Topics by XP")
    #plt.gcf().subplots_adjust(bottom=1)
    plt.savefig('topic_chart.png',bbox_inches="tight")
    plt.show()

    document.add_picture('topic_chart.png')


    
    document.add_page_break()
    document.add_heading('All Courses', level=1)
    print('All Courses')

    course_list = sel.css(css_course_list).extract()
    lang_raw=sel.xpath(xsel_course_lang).extract()
    lang_list=just_words(lang_raw,is_amp,['0'])
    course_urls=sel.css(css_course_urls).extract()
    course_descriptions=just_words(\
                                   sel.css(css_course_descriptions).extract(),\
                                   nothing)

    my_courses = pd.DataFrame(list(zip(lang_list,course_list,course_list,\
                                       course_descriptions,course_urls)), \
                              columns =['Technology','Course_Name','Course_Raw'\
                                        ,'Description','URL'])

    my_courses.Technology = my_courses.Technology.str.title()
    my_courses.Course_Name = my_courses.Course_Name.str.title()
    my_courses['course_lower'] = my_courses.Course_Name.str.lower().str.strip()
    
    # add cert data
    courses_join_df = pd.DataFrame.merge(my_courses, cert_df, how='left', 
                                         on='course_lower', 
                                         sort=False, suffixes=('_x', '_y'))

    my_courses.groupby('Technology', as_index=False)['Course_Name'].count().plot('Technology', kind='bar')
    plt.xticks(rotation=45)
    plt.title("Courses by Technology")
    plt.ylabel('Courses')
    #plt.gcf().subplots_adjust(bottom=1)
    plt.savefig('tech_chart.png',bbox_inches="tight")
    plt.show()
    document.add_picture('tech_chart.png')

    tech_table=my_courses. \
    groupby('Technology',as_index=False)['Course_Name']. \
    count().rename(columns={"Course_Name":"Course_Count"})

    doctable(tech_table,"Course count by Technology (Language)")
    #    row_cells[2].text = desc
    print(tech_table)

    document.add_page_break()
    python_courses = courses_join_df[['ref','Course_Name','date','Technology']]\
    .loc[courses_join_df.Technology=='Python'] \
          .sort_values(by='Course_Name').reset_index(drop=True)
          
    doctable(python_courses[['ref','Course_Name','date']],"Python Courses List")

    print('Python Courses')
    print(python_courses)
    r_courses = my_courses \
        .loc[my_courses.Technology=='R','Course_Name'] \
        .sort_values().reset_index(drop=True)
    
    document.add_page_break()
    
    r_courses = courses_join_df[['ref','Course_Name','date','Technology']]\
    .loc[courses_join_df.Technology=='R'] \
          .sort_values(by='Course_Name').reset_index(drop=True)
          
    doctable(r_courses[['ref','Course_Name','date']],"R Course List")
    print('R Courses')
    print(r_courses)

    document.add_page_break()
    track_names = sel.css(css_track_names).extract()

    track_names = just_words(track_names,is_amp,['0'])
    my_tracks = pd.DataFrame(list(track_names), \
                             columns =['Skill_Track'])
    my_tracks.Skill_Track = my_tracks.Skill_Track.str.\
                            replace('\n ','').\
                            str.strip()
    # add cert data
    tracks_join_df = pd.DataFrame.merge(my_tracks, cert_df, how='left', 
                                        left_on='Skill_Track', right_on='course', 
                                        sort=False, suffixes=('_x', '_y'))

    doctable(tracks_join_df[['ref','course','c_type']],"Career and Skill Tracks")
    
    document.add_page_break()
    print("Skill Tracks")
    print(my_tracks)
    
    project_names = sel.css(css_project_names).extract()
    project_urls = sel.css(css_project_urls).extract()
    my_projects = pd.DataFrame(list(zip(project_names,project_urls)), \
                               columns =['Project','URL'])

    doctable(my_projects.Project,"Projects")
    print("Projects")
    print(my_projects.Project)

    my_tech = my_courses.Technology.unique()

    document.save('test.docx')
